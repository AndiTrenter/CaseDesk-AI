"""
CaseDesk AI - Response Generator Service
AI-powered response generation with PDF/DOCX output and document bundling
"""
import os
import json
import smtplib
import zipfile
import tempfile
import re
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime, timezone
from typing import Optional, List, Dict, Any
from pathlib import Path
import logging
import uuid
import aiofiles

logger = logging.getLogger(__name__)


class ResponseGeneratorService:
    """AI-powered response generation with PDF/DOCX output and document bundling"""
    
    def __init__(self, db, ai_service):
        self.db = db
        self.ai_service = ai_service
    
    async def analyze_case_requirements(self, case_id: str, user_id: str) -> Dict[str, Any]:
        """Analyze case and determine required documents and response type"""
        case = await self.db.cases.find_one(
            {"id": case_id, "user_id": user_id}, {"_id": 0}
        )
        if not case:
            return {"success": False, "error": "Fall nicht gefunden"}
        
        documents = await self.db.documents.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).to_list(100)
        
        emails = await self.db.emails.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).to_list(100)
        
        history = await self.db.correspondence.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).sort("created_at", -1).to_list(50)
        
        context = f"""
FALL: {case.get('title')}
BESCHREIBUNG: {case.get('description', 'Keine')}
AKTENZEICHEN: {case.get('reference_number', 'Keines')}
STATUS: {case.get('status')}

VERKNUEPFTE DOKUMENTE ({len(documents)}):
"""
        for doc in documents:
            context += f"\n- {doc.get('display_name', doc.get('original_filename'))}"
            context += f"\n  Typ: {doc.get('document_type')}, Absender: {doc.get('sender', 'Unbekannt')}"
            if doc.get('ai_summary'):
                context += f"\n  Zusammenfassung: {doc.get('ai_summary')}"
            if doc.get('ocr_text'):
                context += f"\n  Inhalt (Auszug): {doc.get('ocr_text', '')[:500]}..."
        
        if emails:
            context += f"\n\nVERKNUEPFTE E-MAILS ({len(emails)}):"
            for em in emails[:5]:
                context += f"\n- {em.get('subject')} von {em.get('sender')}"
                if em.get('ai_summary'):
                    context += f"\n  Zusammenfassung: {em.get('ai_summary')}"
        
        if history:
            context += f"\n\nBISHERIGE KORRESPONDENZ ({len(history)}):"
            for h in history[:5]:
                context += f"\n- {h.get('created_at', '')[:10]}: {h.get('subject')}"
                context += f" ({h.get('status', 'erstellt')})"
        
        system_prompt = """Du bist ein Experte fuer Fallanalyse und Korrespondenz.
Analysiere den Fall und identifiziere:
1. Art der erforderlichen Antwort
2. Welche Dokumente/Nachweise benoetigt werden
3. Fristen die beachtet werden muessen
4. Empfohlene Vorgehensweise

Antworte NUR mit einem validen JSON:
{
    "antworttyp": "z.B. Widerspruch, Antrag, Stellungnahme",
    "empfaenger": "Name/Behoerde an die geantwortet werden soll",
    "benoetigt_dokumente": ["Liste der benoetigten Dokumente"],
    "verfuegbare_dokumente": ["Liste der bereits vorhandenen passenden Dokumente"],
    "fehlende_dokumente": ["Liste der noch zu beschaffenden Dokumente"],
    "fristen": ["Erkannte Fristen mit Datum"],
    "empfehlung": "Kurze Empfehlung zur Vorgehensweise",
    "dringlichkeit": "hoch/mittel/niedrig"
}"""

        prompt = f"""Analysiere diesen Fall und identifiziere die Anforderungen:\n\n{context}"""

        try:
            response = await self.ai_service.generate(prompt, system_prompt)
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                analysis = json.loads(json_match.group())
                return {
                    "success": True, "case": case, "documents": documents,
                    "emails": emails, "history": history, "analysis": analysis
                }
            else:
                return {
                    "success": True, "case": case, "documents": documents,
                    "emails": emails, "history": history,
                    "analysis": {"raw_response": response}
                }
        except Exception as e:
            logger.error(f"Case analysis error: {e}")
            return {
                "success": True, "case": case, "documents": documents,
                "emails": emails, "history": history,
                "analysis": None, "error": str(e)
            }
    
    async def generate_response(
        self, case_id: str, user_id: str, response_type: str,
        recipient: str, subject: str, instructions: str = None,
        include_document_ids: List[str] = None,
        output_format: str = "pdf"
    ) -> Dict[str, Any]:
        """Generate a complete response package with PDF or DOCX output"""
        
        case = await self.db.cases.find_one(
            {"id": case_id, "user_id": user_id}, {"_id": 0}
        )
        if not case:
            return {"success": False, "error": "Fall nicht gefunden"}
        
        user = await self.db.users.find_one(
            {"id": user_id}, {"_id": 0, "password_hash": 0}
        )
        
        documents = []
        if include_document_ids:
            documents = await self.db.documents.find(
                {"id": {"$in": include_document_ids}, "user_id": user_id},
                {"_id": 0}
            ).to_list(100)
        
        all_case_docs = await self.db.documents.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).to_list(100)
        
        context = f"""
FALL: {case.get('title')}
AKTENZEICHEN: {case.get('reference_number', 'nicht angegeben')}
BESCHREIBUNG: {case.get('description', '')}

ABSENDER (Nutzer):
Name: {user.get('full_name', user.get('username'))}
E-Mail: {user.get('email')}

EMPFAENGER: {recipient}
BETREFF: {subject}
ART DES SCHREIBENS: {response_type}

VERFUEGBARE DOKUMENTE IM FALL:
"""
        for doc in all_case_docs:
            context += f"\n- {doc.get('display_name', doc.get('original_filename'))}"
            if doc.get('ai_summary'):
                context += f": {doc.get('ai_summary')}"
            if doc.get('ocr_text'):
                context += f"\n  Inhalt: {doc.get('ocr_text', '')[:300]}..."
        
        if include_document_ids:
            context += "\n\nALS ANLAGE BEIZUFUEGEN:"
            for doc in documents:
                context += f"\n- {doc.get('display_name', doc.get('original_filename'))}"
        
        system_prompt = """Du bist ein Experte fuer formelle Korrespondenz.
Erstelle ein professionelles, korrektes Antwortschreiben.

WICHTIGE REGELN:
- Schreibe formal und hoeflich
- Verwende das korrekte Format fuer offizielle Schreiben
- Beziehe dich auf vorhandene Aktenzeichen
- Erwaehne beigefuegte Anlagen
- NIEMALS falsche Angaben oder erfundene Fakten
- Nur auf Basis der vorhandenen Dokumente argumentieren

FORMAT:
- Absender oben rechts
- Empfaenger links
- Datum
- Betreff mit Aktenzeichen
- Anrede
- Text
- Grussformel
- Unterschrift
- Anlagen-Verzeichnis"""

        prompt = f"""Erstelle ein {response_type} basierend auf folgendem Kontext:

{context}

ZUSAETZLICHE ANWEISUNGEN: {instructions or 'Keine weiteren Anweisungen'}

Erstelle das vollstaendige Schreiben."""

        try:
            response_text = await self.ai_service.generate(prompt, system_prompt, max_tokens=3000)
            
            # Generate file in requested format
            file_path = None
            file_name = None
            if output_format == "docx":
                file_path, file_name = self._generate_docx(response_text, subject, user)
            elif output_format == "pdf":
                file_path, file_name = self._generate_pdf(response_text, subject, user)
            
            correspondence_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()
            
            correspondence = {
                "id": correspondence_id,
                "user_id": user_id,
                "case_id": case_id,
                "type": response_type,
                "subject": subject,
                "recipient": recipient,
                "content": response_text,
                "document_ids": include_document_ids or [],
                "output_format": output_format,
                "file_path": file_path,
                "file_name": file_name,
                "status": "draft",
                "created_at": now,
                "updated_at": now,
                "sent_at": None,
                "sent_via": None
            }
            
            await self.db.correspondence.insert_one(correspondence)
            
            return {
                "success": True,
                "correspondence_id": correspondence_id,
                "content": response_text,
                "output_format": output_format,
                "file_name": file_name,
                "documents": documents,
                "case": case
            }
            
        except Exception as e:
            logger.error(f"Response generation error: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_pdf(self, text: str, subject: str, user: dict) -> tuple:
        """Generate a PDF file from the response text"""
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=25)
            
            # Use built-in font with Latin-1 encoding fallback
            pdf.set_font("Helvetica", size=11)
            
            # Header
            pdf.set_font("Helvetica", "B", 14)
            safe_subject = subject.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 10, safe_subject, ln=True, align='C')
            pdf.ln(5)
            
            # Date
            pdf.set_font("Helvetica", size=9)
            pdf.cell(0, 6, f"Erstellt am: {datetime.now().strftime('%d.%m.%Y')}", ln=True, align='R')
            pdf.ln(5)
            
            # Body
            pdf.set_font("Helvetica", size=11)
            for line in text.split('\n'):
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, safe_line)
            
            # Save
            temp_dir = tempfile.mkdtemp()
            clean_subject = re.sub(r'[^\w\s-]', '', subject)[:40]
            file_name = f"Schreiben_{clean_subject}.pdf"
            file_path = os.path.join(temp_dir, file_name)
            pdf.output(file_path)
            
            return file_path, file_name
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return None, None
    
    def _generate_docx(self, text: str, subject: str, user: dict) -> tuple:
        """Generate a DOCX file from the response text"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            style = doc.styles['Normal']
            style.font.size = Pt(11)
            style.font.name = 'Calibri'
            
            # Title
            title = doc.add_heading(subject, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"Erstellt am: {datetime.now().strftime('%d.%m.%Y')}")
            date_run.font.size = Pt(9)
            
            doc.add_paragraph()
            
            # Body
            for line in text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
            
            # Save
            temp_dir = tempfile.mkdtemp()
            clean_subject = re.sub(r'[^\w\s-]', '', subject)[:40]
            file_name = f"Schreiben_{clean_subject}.docx"
            file_path = os.path.join(temp_dir, file_name)
            doc.save(file_path)
            
            return file_path, file_name
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return None, None
    
    async def create_download_package(
        self, correspondence_id: str, user_id: str
    ) -> Dict[str, Any]:
        """Create a ZIP package with response letter (PDF/DOCX) and attachments"""
        
        correspondence = await self.db.correspondence.find_one(
            {"id": correspondence_id, "user_id": user_id}, {"_id": 0}
        )
        if not correspondence:
            return {"success": False, "error": "Korrespondenz nicht gefunden"}
        
        documents = []
        if correspondence.get("document_ids"):
            documents = await self.db.documents.find(
                {"id": {"$in": correspondence["document_ids"]}, "user_id": user_id},
                {"_id": 0}
            ).to_list(100)
        
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, f"antwort_{correspondence_id[:8]}.zip")
        
        try:
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Add the generated file (PDF/DOCX) if available
                if correspondence.get("file_path") and os.path.exists(correspondence["file_path"]):
                    zipf.write(correspondence["file_path"], correspondence.get("file_name", "Schreiben.pdf"))
                else:
                    # Fallback: add as text file
                    letter_filename = f"Schreiben_{correspondence.get('subject', 'Antwort')[:30]}.txt"
                    letter_filename = letter_filename.replace("/", "-").replace("\\", "-")
                    zipf.writestr(letter_filename, correspondence.get("content", ""))
                
                # Add attached documents
                for i, doc in enumerate(documents, 1):
                    if os.path.exists(doc.get("storage_path", "")):
                        ext = Path(doc["storage_path"]).suffix
                        doc_filename = f"Anlage_{i}_{doc.get('display_name', doc.get('original_filename', 'dokument'))}"
                        if not doc_filename.endswith(ext):
                            doc_filename += ext
                        doc_filename = doc_filename.replace("/", "-").replace("\\", "-")
                        zipf.write(doc["storage_path"], doc_filename)
            
            return {
                "success": True,
                "zip_path": zip_path,
                "filename": f"antwort_{correspondence_id[:8]}.zip"
            }
            
        except Exception as e:
            logger.error(f"Package creation error: {e}")
            return {"success": False, "error": str(e)}
    
    async def send_via_email(
        self, correspondence_id: str, user_id: str,
        mail_account_id: str, recipient_email: str
    ) -> Dict[str, Any]:
        """Send correspondence via email with attachments"""
        
        correspondence = await self.db.correspondence.find_one(
            {"id": correspondence_id, "user_id": user_id}, {"_id": 0}
        )
        if not correspondence:
            return {"success": False, "error": "Korrespondenz nicht gefunden"}
        
        account = await self.db.mail_accounts.find_one(
            {"id": mail_account_id, "user_id": user_id}, {"_id": 0}
        )
        if not account:
            return {"success": False, "error": "E-Mail-Konto nicht gefunden"}
        
        documents = []
        if correspondence.get("document_ids"):
            documents = await self.db.documents.find(
                {"id": {"$in": correspondence["document_ids"]}, "user_id": user_id},
                {"_id": 0}
            ).to_list(100)
        
        try:
            msg = MIMEMultipart()
            msg['From'] = account['email']
            msg['To'] = recipient_email
            msg['Subject'] = correspondence.get('subject', 'Antwort')
            
            msg.attach(MIMEText(correspondence.get('content', ''), 'plain', 'utf-8'))
            
            # Attach the generated letter file (PDF/DOCX)
            if correspondence.get("file_path") and os.path.exists(correspondence["file_path"]):
                with open(correspondence["file_path"], "rb") as f:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{correspondence.get("file_name", "Schreiben")}"')
                msg.attach(part)
            
            for doc in documents:
                if os.path.exists(doc.get("storage_path", "")):
                    with open(doc["storage_path"], "rb") as f:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                    encoders.encode_base64(part)
                    filename = doc.get('display_name', doc.get('original_filename', 'dokument'))
                    part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
                    msg.attach(part)
            
            smtp_server = account.get('smtp_server', account.get('imap_server', '').replace('imap', 'smtp'))
            smtp_port = account.get('smtp_port', 587)
            
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                if account.get('smtp_use_tls', True):
                    server.starttls()
                server.login(account['email'], account['password'])
                server.send_message(msg)
            
            now = datetime.now(timezone.utc).isoformat()
            await self.db.correspondence.update_one(
                {"id": correspondence_id},
                {"$set": {"status": "sent", "sent_at": now, "sent_via": "email", "sent_to": recipient_email, "updated_at": now}}
            )
            
            return {"success": True, "message": f"E-Mail erfolgreich an {recipient_email} gesendet"}
            
        except smtplib.SMTPAuthenticationError:
            return {"success": False, "error": "SMTP-Authentifizierung fehlgeschlagen"}
        except smtplib.SMTPException as e:
            return {"success": False, "error": f"SMTP-Fehler: {str(e)}"}
        except Exception as e:
            logger.error(f"Email send error: {e}")
            return {"success": False, "error": str(e)}
    
    async def get_case_correspondence_history(self, case_id: str, user_id: str) -> List[Dict[str, Any]]:
        """Get all correspondence for a case"""
        history = await self.db.correspondence.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).sort("created_at", -1).to_list(100)
        return history
    
    async def update_correspondence(
        self, correspondence_id: str, user_id: str,
        content: str = None, subject: str = None, status: str = None
    ) -> Dict[str, Any]:
        """Update correspondence"""
        update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
        if content is not None:
            update_data["content"] = content
        if subject is not None:
            update_data["subject"] = subject
        if status is not None:
            update_data["status"] = status
        
        result = await self.db.correspondence.update_one(
            {"id": correspondence_id, "user_id": user_id},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            return {"success": True}
        return {"success": False, "error": "Nicht gefunden oder keine Aenderungen"}
    
    async def delete_correspondence(self, correspondence_id: str, user_id: str) -> Dict[str, Any]:
        """Delete correspondence"""
        result = await self.db.correspondence.delete_one(
            {"id": correspondence_id, "user_id": user_id}
        )
        if result.deleted_count > 0:
            return {"success": True}
        return {"success": False, "error": "Nicht gefunden"}
