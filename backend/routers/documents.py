"""Documents Router"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Query
from fastapi.responses import FileResponse
from datetime import datetime, timezone
from typing import List, Optional
from pathlib import Path
import uuid
import os
import json
import httpx
import aiofiles
import logging
import io
import zipfile
import xml.etree.ElementTree as ET

from deps import db, require_auth, log_action, UPLOAD_DIR, OCR_SERVICE_URL, create_download_token, verify_download_token
from models import Document

logger = logging.getLogger(__name__)
router = APIRouter()


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from .docx (Word 2007+) files"""
    try:
        # .docx is a ZIP file containing XML
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            # Main document content is in word/document.xml
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml')
                tree = ET.fromstring(xml_content)
                
                # Extract all text from paragraphs
                # Word XML namespace
                word_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                
                text_parts = []
                for para in tree.iter(f'{word_ns}p'):
                    para_text = []
                    for text_elem in para.iter(f'{word_ns}t'):
                        if text_elem.text:
                            para_text.append(text_elem.text)
                    if para_text:
                        text_parts.append(''.join(para_text))
                
                return '\n'.join(text_parts)
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
    return ""


def extract_text_from_odt(content: bytes) -> str:
    """Extract text from .odt (OpenDocument) files"""
    try:
        # .odt is a ZIP file containing XML
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            # Main document content is in content.xml
            if 'content.xml' in zf.namelist():
                xml_content = zf.read('content.xml')
                tree = ET.fromstring(xml_content)
                
                # Extract all text elements
                text_parts = []
                for elem in tree.iter():
                    if elem.text and elem.text.strip():
                        text_parts.append(elem.text.strip())
                    if elem.tail and elem.tail.strip():
                        text_parts.append(elem.tail.strip())
                
                return ' '.join(text_parts)
    except Exception as e:
        logger.warning(f"ODT extraction failed: {e}")
    return ""


def extract_text_from_rtf(content: bytes) -> str:
    """Extract text from .rtf (Rich Text Format) files"""
    try:
        text = content.decode('latin-1', errors='ignore')
        # Simple RTF text extraction - remove control words
        import re
        # Remove RTF control words
        text = re.sub(r'\\[a-z]+\d*\s?', '', text)
        # Remove braces
        text = re.sub(r'[{}]', '', text)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    except Exception as e:
        logger.warning(f"RTF extraction failed: {e}")
    return ""


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from plain text files"""
    try:
        # Try UTF-8 first, then fallback encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
    except Exception as e:
        logger.warning(f"TXT extraction failed: {e}")
    return ""



def extract_text_from_xlsx(content: bytes) -> str:
    """Extract text from .xlsx (Excel 2007+) files"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), data_only=True)
        
        text_parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Tabelle: {sheet_name} ===")
            
            for row in sheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"XLSX extraction failed: {e}")
    return ""


def extract_text_from_xls(content: bytes) -> str:
    """Extract text from .xls (Excel 97-2003) files"""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=content)
        
        text_parts = []
        for sheet_name in wb.sheet_names():
            sheet = wb.sheet_by_name(sheet_name)
            text_parts.append(f"=== Tabelle: {sheet_name} ===")
            
            for row_idx in range(sheet.nrows):
                row_text = []
                for col_idx in range(sheet.ncols):
                    cell_value = sheet.cell_value(row_idx, col_idx)
                    if cell_value:
                        row_text.append(str(cell_value))
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"XLS extraction failed: {e}")
    return ""


def extract_text_from_ods(content: bytes) -> str:
    """Extract text from .ods (OpenDocument Spreadsheet) files"""
    try:
        from odf import text as odf_text
        from odf.opendocument import load
        
        doc = load(io.BytesIO(content))
        text_parts = []
        
        # Extract all text from the spreadsheet
        for element in doc.getElementsByType(odf_text.P):
            if element.firstChild:
                text_parts.append(str(element))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"ODS extraction failed: {e}")
    return ""

def extract_text_from_pdf(content: bytes) -> str:
    """Fallback: Extract text directly from PDF, with OCR for scanned pages"""
    text_parts = []

    # First try: Extract embedded text with PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")

    # If we got text from embedded text, return it
    if text_parts and sum(len(t) for t in text_parts) > 50:
        return "\n".join(text_parts)

    # Second try: OCR scanned pages with tesseract
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(content, dpi=300)
        ocr_parts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang='deu+eng')
            if page_text and page_text.strip():
                ocr_parts.append(page_text)
        if ocr_parts:
            return "\n".join(ocr_parts)
    except Exception as e:
        logger.warning(f"Tesseract OCR fallback failed: {e}")

    return "\n".join(text_parts)


async def try_ocr_or_fallback(filename: str, content: bytes, mime_type: str) -> str:
    """Try OCR service first, fall back to direct text extraction for various formats"""
    ocr_text = ""
    file_ext = os.path.splitext(filename.lower())[1]
    
    # Handle Word documents (.docx)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_ext == ".docx":
        logger.info(f"Extracting text from DOCX: {filename}")
        ocr_text = extract_text_from_docx(content)
        if ocr_text:
            return ocr_text
    
    # Handle older Word documents (.doc) - try python-docx fallback
    if mime_type == "application/msword" or file_ext == ".doc":
        logger.info(f"Extracting text from DOC: {filename}")
        # .doc files are binary, try to extract what we can
        try:
            # First try: if it's actually a .docx renamed to .doc
            ocr_text = extract_text_from_docx(content)
            if ocr_text:
                return ocr_text
        except Exception:
            pass
        # Fallback: extract readable strings
        try:
            text = content.decode('latin-1', errors='ignore')
            import re
            # Extract readable text strings
            readable = re.findall(r'[\x20-\x7E\xC0-\xFF]{10,}', text)
            ocr_text = ' '.join(readable)
            if ocr_text:
                return ocr_text
        except Exception as e:
            logger.warning(f"DOC extraction failed: {e}")
    
    # Handle OpenDocument (.odt)
    if mime_type == "application/vnd.oasis.opendocument.text" or file_ext == ".odt":
        logger.info(f"Extracting text from ODT: {filename}")
        ocr_text = extract_text_from_odt(content)
        if ocr_text:
            return ocr_text
    
    # Handle RTF files
    if mime_type == "application/rtf" or mime_type == "text/rtf" or file_ext == ".rtf":
        logger.info(f"Extracting text from RTF: {filename}")
        ocr_text = extract_text_from_rtf(content)
        if ocr_text:
            return ocr_text
    
    # Handle plain text files
    if mime_type.startswith("text/") or file_ext in [".txt", ".csv", ".md", ".log"]:
        logger.info(f"Extracting text from TXT: {filename}")
        ocr_text = extract_text_from_txt(content)
        if ocr_text:
            return ocr_text
    
    # Handle Excel files (.xlsx)
    if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" or file_ext == ".xlsx":
        logger.info(f"Extracting text from XLSX: {filename}")
        ocr_text = extract_text_from_xlsx(content)
        if ocr_text:
            return ocr_text
    
    # Handle older Excel files (.xls)
    if mime_type == "application/vnd.ms-excel" or file_ext == ".xls":
        logger.info(f"Extracting text from XLS: {filename}")
        ocr_text = extract_text_from_xls(content)
        if ocr_text:
            return ocr_text
    
    # Handle OpenDocument Spreadsheet (.ods)
    if mime_type == "application/vnd.oasis.opendocument.spreadsheet" or file_ext == ".ods":
        logger.info(f"Extracting text from ODS: {filename}")
        ocr_text = extract_text_from_ods(content)
        if ocr_text:
            return ocr_text

    # Try OCR service for PDFs and images
    if mime_type in ["application/pdf", "image/png", "image/jpeg", "image/tiff", "image/gif", "image/webp"]:
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                ocr_response = await client.post(
                    f"{OCR_SERVICE_URL}/ocr",
                    files={"file": (filename, content, mime_type)}
                )
                if ocr_response.status_code == 200:
                    ocr_data = ocr_response.json()
                    ocr_text = ocr_data.get("text", "")
        except Exception as e:
            logger.warning(f"OCR service unavailable: {e}")

    # Fallback: direct PDF text extraction
    if not ocr_text and mime_type == "application/pdf":
        logger.info("Using fallback PDF text extraction")
        ocr_text = extract_text_from_pdf(content)

    return ocr_text


@router.get("/documents")
async def list_documents(
    case_id: Optional[str] = None,
    search: Optional[str] = None,
    document_type: Optional[str] = None,
    unassigned: Optional[bool] = None,
    user: dict = Depends(require_auth)
):
    query = {"user_id": user["id"]}
    if case_id:
        query["case_id"] = case_id
    if document_type:
        query["document_type"] = document_type
    if unassigned:
        query["case_id"] = None
    
    if search:
        try:
            text_results = await db.documents.find(
                {"$text": {"$search": search}, "user_id": user["id"]},
                {"_id": 0, "score": {"$meta": "textScore"}}
            ).sort([("score", {"$meta": "textScore"})]).to_list(100)
            
            if text_results:
                return text_results
        except Exception:
            pass
        
        query["$or"] = [
            {"display_name": {"$regex": search, "$options": "i"}},
            {"original_filename": {"$regex": search, "$options": "i"}},
            {"tags": {"$regex": search, "$options": "i"}},
            {"ai_summary": {"$regex": search, "$options": "i"}}
        ]
    
    documents = await db.documents.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return documents


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    case_id: str = Form(None),
    document_type: str = Form("other"),
    user: dict = Depends(require_auth)
):
    now = datetime.now(timezone.utc).isoformat()
    doc_id = str(uuid.uuid4())
    
    user_upload_dir = UPLOAD_DIR / user["id"]
    user_upload_dir.mkdir(parents=True, exist_ok=True)
    
    file_ext = Path(file.filename).suffix
    storage_filename = f"{doc_id}{file_ext}"
    storage_path = user_upload_dir / storage_filename
    
    content = await file.read()
    async with aiofiles.open(storage_path, 'wb') as f:
        await f.write(content)
    
    document = {
        "id": doc_id,
        "user_id": user["id"],
        "case_id": case_id,
        "filename": file.filename,
        "original_filename": file.filename,
        "display_name": file.filename,
        "storage_path": str(storage_path),
        "mime_type": file.content_type or "application/octet-stream",
        "size": len(content),
        "document_type": document_type,
        "ocr_text": None,
        "ocr_processed": False,
        "ai_analyzed": False,
        "tags": [],
        "metadata": {},
        "created_at": now,
        "updated_at": now
    }
    
    await db.documents.insert_one(document)
    
    if case_id:
        await db.cases.update_one(
            {"id": case_id, "user_id": user["id"]},
            {"$addToSet": {"document_ids": doc_id}}
        )
    
    await log_action(user["id"], "upload_document", "document", doc_id, {"filename": file.filename})
    
    # Background AI processing
    try:
        ocr_text = await try_ocr_or_fallback(file.filename, content, file.content_type)
        
        if ocr_text:
            from ai_service import get_ai_service, DocumentAnalyzer
            ai_service = await get_ai_service(db)
            analyzer = DocumentAnalyzer(ai_service)
            
            analysis = await analyzer.analyze_document(ocr_text, file.filename)
            
            if analysis.get("success"):
                metadata = analysis.get("metadata", {})
                new_display_name = analysis.get("new_filename", file.filename)
                
                update_data = {
                    "ocr_text": ocr_text,
                    "ocr_processed": True,
                    "ai_analyzed": True,
                    "display_name": new_display_name,
                    "document_type": metadata.get("dokumenttyp", document_type).lower().replace("ae", "ae").replace("oe", "oe").replace("ue", "ue"),
                    "tags": metadata.get("tags", []),
                    "ai_summary": metadata.get("zusammenfassung"),
                    "sender": metadata.get("absender"),
                    "importance": metadata.get("wichtigkeit", "mittel"),
                    "deadlines": metadata.get("fristen", []),
                    "metadata": metadata,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }
                
                if metadata.get("datum") and metadata["datum"] != "null":
                    try:
                        update_data["document_date"] = metadata["datum"]
                    except Exception:
                        pass
                
                await db.documents.update_one({"id": doc_id}, {"$set": update_data})
                document.update(update_data)
                
                # Create tasks AND calendar events for deadlines
                for deadline in metadata.get("fristen", []):
                    try:
                        task = {
                            "id": str(uuid.uuid4()),
                            "user_id": user["id"],
                            "title": f"Frist: {metadata.get('kurzthema', 'Dokument')}",
                            "description": f"Automatisch erkannte Frist aus: {new_display_name}",
                            "priority": "high",
                            "status": "todo",
                            "due_date": deadline if isinstance(deadline, str) else deadline.get("datum"),
                            "case_id": case_id,
                            "document_id": doc_id,
                            "source": "document_ai",
                            "created_at": now,
                            "updated_at": now
                        }
                        await db.tasks.insert_one(task)
                    except Exception:
                        pass
                
                # Auto-create calendar events from deadlines
                from routers.events import create_events_from_deadlines
                await create_events_from_deadlines(
                    user["id"], metadata.get("fristen", []),
                    new_display_name, case_id, doc_id
                )
            else:
                await db.documents.update_one(
                    {"id": doc_id},
                    {"$set": {"ocr_text": ocr_text, "ocr_processed": True, "updated_at": now}}
                )
    except Exception as e:
        logger.error(f"Document processing error: {e}")
    
    document.pop("_id", None)
    return {"success": True, "document": document}


@router.post("/documents/{document_id}/reprocess")
async def reprocess_document(document_id: str, force: bool = False, user: dict = Depends(require_auth)):
    """Reprocess document - extracts text and analyzes with AI
    
    Args:
        force: If True, reprocess even if ocr_text already exists
    """
    document = await db.documents.find_one(
        {"id": document_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    ocr_text = document.get("ocr_text")
    
    # Reprocess if no text OR if force flag is set
    if (not ocr_text or force) and document.get("storage_path") and os.path.exists(document["storage_path"]):
        try:
            async with aiofiles.open(document["storage_path"], 'rb') as f:
                content = await f.read()
            
            ocr_text = await try_ocr_or_fallback(
                document["original_filename"], content, document["mime_type"]
            )
            logger.info(f"Extracted {len(ocr_text) if ocr_text else 0} chars from {document['original_filename']}")
        except Exception as e:
            logger.warning(f"OCR reprocess error: {e}")
    
    if ocr_text:
        # Always save the extracted text first
        update_data = {
            "ocr_text": ocr_text,
            "ocr_processed": True,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        
        # Try AI analysis (may fail if AI service not available)
        try:
            from ai_service import get_ai_service, DocumentAnalyzer
            ai_service = await get_ai_service(db)
            analyzer = DocumentAnalyzer(ai_service)
            
            analysis = await analyzer.analyze_document(ocr_text, document["original_filename"])
            
            if analysis.get("success"):
                metadata = analysis.get("metadata", {})
                new_display_name = analysis.get("new_filename", document["display_name"])
                
                update_data.update({
                    "ai_analyzed": True,
                    "display_name": new_display_name,
                    "tags": metadata.get("tags", []),
                    "ai_summary": metadata.get("zusammenfassung"),
                    "sender": metadata.get("absender"),
                    "deadlines": metadata.get("fristen", []),
                    "metadata": metadata,
                })
                
                if metadata.get("datum") and metadata["datum"] != "null":
                    update_data["document_date"] = metadata["datum"]
                
                # Create events from deadlines
                from routers.events import create_events_from_deadlines
                await create_events_from_deadlines(
                    user["id"], metadata.get("fristen", []),
                    new_display_name, document.get("case_id"), document_id
                )
        except Exception as e:
            logger.warning(f"AI analysis failed, but text extraction succeeded: {e}")
        
        # Update document with extracted text (and AI analysis if successful)
        await db.documents.update_one({"id": document_id}, {"$set": update_data})
        
        return {
            "success": True, 
            "message": "Document reprocessed", 
            "display_name": update_data.get("display_name", document["display_name"]),
            "text_extracted": len(ocr_text),
            "ai_analyzed": update_data.get("ai_analyzed", False)
        }
    
    return {"success": False, "error": "No text to process"}


@router.post("/documents/{document_id}/ocr")
async def process_document_ocr(document_id: str, user: dict = Depends(require_auth)):
    document = await db.documents.find_one(
        {"id": document_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.get("storage_path", "")):
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        async with aiofiles.open(document["storage_path"], 'rb') as f:
            content = await f.read()
        
        ocr_text = await try_ocr_or_fallback(
            document["original_filename"], content, document["mime_type"]
        )
        
        if ocr_text:
            await db.documents.update_one(
                {"id": document_id},
                {"$set": {"ocr_text": ocr_text, "ocr_processed": True, "updated_at": datetime.now(timezone.utc).isoformat()}}
            )
            return {"success": True, "ocr_text": ocr_text[:500]}
        else:
            return {"success": False, "error": "No text could be extracted"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.get("/documents/{document_id}", response_model=Document)
async def get_document(document_id: str, user: dict = Depends(require_auth)):
    document = await db.documents.find_one({"id": document_id, "user_id": user["id"]}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return Document(**document)


@router.delete("/documents/{document_id}")
async def delete_document(document_id: str, user: dict = Depends(require_auth)):
    document = await db.documents.find_one({"id": document_id, "user_id": user["id"]})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document.get("storage_path") and os.path.exists(document["storage_path"]):
        os.remove(document["storage_path"])
    
    if document.get("case_id"):
        await db.cases.update_one(
            {"id": document["case_id"]},
            {"$pull": {"document_ids": document_id}}
        )
    
    await db.documents.delete_one({"id": document_id})
    await log_action(user["id"], "delete_document", "document", document_id)
    return {"success": True, "message": "Document deleted"}


@router.get("/documents/{document_id}/preview")
async def get_document_preview(document_id: str, user: dict = Depends(require_auth)):
    document = await db.documents.find_one(
        {"id": document_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": document["id"],
        "display_name": document.get("display_name", document["original_filename"]),
        "mime_type": document["mime_type"],
        "size": document["size"],
        "ocr_text": document.get("ocr_text"),
        "ai_summary": document.get("ai_summary"),
        "tags": document.get("tags", []),
        "metadata": document.get("metadata", {}),
        "sender": document.get("sender"),
        "document_date": document.get("document_date"),
        "deadlines": document.get("deadlines", []),
        "importance": document.get("importance")
    }


@router.get("/documents/{document_id}/download")
async def download_document(document_id: str, user: dict = Depends(require_auth)):
    document = await db.documents.find_one(
        {"id": document_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = document["storage_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        file_path,
        filename=document.get("display_name", document["original_filename"]),
        media_type=document["mime_type"]
    )


@router.get("/documents/{document_id}/download-token")
async def get_download_token(document_id: str, user: dict = Depends(require_auth)):
    """Generate a short-lived token for downloading the document"""
    document = await db.documents.find_one(
        {"id": document_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    token = create_download_token(document_id, user["id"], expires_minutes=5)
    return {"token": token, "expires_in": 300}


@router.get("/documents/{document_id}/view")
async def view_document_with_token(document_id: str, token: str = Query(...)):
    """View/download document using a token (no auth header needed)"""
    payload = verify_download_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    if payload.get("doc_id") != document_id:
        raise HTTPException(status_code=403, detail="Token does not match document")
    
    document = await db.documents.find_one(
        {"id": document_id, "user_id": payload["user_id"]}, {"_id": 0}
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = document["storage_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Return file for inline viewing (Content-Disposition: inline)
    from fastapi.responses import Response
    async with aiofiles.open(file_path, 'rb') as f:
        content = await f.read()
    
    return Response(
        content=content,
        media_type=document["mime_type"],
        headers={
            "Content-Disposition": f'inline; filename="{document.get("display_name", document["original_filename"])}"'
        }
    )


@router.get("/documents/{document_id}/auto-link")
async def auto_link_document(document_id: str, user: dict = Depends(require_auth)):
    from ai_service import get_ai_service, ProactiveAssistant
    ai_service = await get_ai_service(db)
    assistant = ProactiveAssistant(ai_service, db)
    return await assistant.auto_link_documents(user["id"], document_id)


@router.put("/documents/{document_id}")
async def update_document(
    document_id: str,
    display_name: str = Form(None),
    document_type: str = Form(None),
    tags: str = Form(None),
    case_id: str = Form(None),
    user: dict = Depends(require_auth)
):
    document = await db.documents.find_one({"id": document_id, "user_id": user["id"]})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    
    if display_name is not None:
        update_data["display_name"] = display_name
    if document_type is not None:
        update_data["document_type"] = document_type
    if tags is not None:
        update_data["tags"] = json.loads(tags) if tags else []
    
    old_case_id = document.get("case_id")
    if case_id is not None and case_id != old_case_id:
        update_data["case_id"] = case_id if case_id else None
        if old_case_id:
            await db.cases.update_one({"id": old_case_id}, {"$pull": {"document_ids": document_id}})
        if case_id:
            await db.cases.update_one(
                {"id": case_id, "user_id": user["id"]},
                {"$addToSet": {"document_ids": document_id}}
            )
    
    await db.documents.update_one({"id": document_id}, {"$set": update_data})
    await log_action(user["id"], "update_document", "document", document_id)
    
    updated = await db.documents.find_one({"id": document_id}, {"_id": 0})
    return {"success": True, "document": updated}


@router.post("/documents/batch-reprocess")
async def batch_reprocess_documents(user: dict = Depends(require_auth)):
    """Reprocess all unprocessed documents using fallback text extraction + AI analysis"""
    unprocessed = await db.documents.find(
        {"user_id": user["id"], "ocr_processed": {"$ne": True}},
        {"_id": 0}
    ).to_list(100)

    processed = 0
    errors = 0

    for doc in unprocessed:
        try:
            storage_path = doc.get("storage_path", "")
            if not storage_path or not os.path.exists(storage_path):
                continue

            async with aiofiles.open(storage_path, 'rb') as f:
                content = await f.read()

            ocr_text = await try_ocr_or_fallback(
                doc["original_filename"], content, doc["mime_type"]
            )

            if ocr_text:
                from ai_service import get_ai_service, DocumentAnalyzer
                ai_service = await get_ai_service(db)
                analyzer = DocumentAnalyzer(ai_service)
                analysis = await analyzer.analyze_document(ocr_text, doc["original_filename"])

                now = datetime.now(timezone.utc).isoformat()
                if analysis.get("success"):
                    metadata = analysis.get("metadata", {})
                    update_data = {
                        "ocr_text": ocr_text,
                        "ocr_processed": True,
                        "ai_analyzed": True,
                        "display_name": analysis.get("new_filename", doc["display_name"]),
                        "document_type": metadata.get("dokumenttyp", doc.get("document_type", "other")).lower(),
                        "tags": metadata.get("tags", []),
                        "ai_summary": metadata.get("zusammenfassung"),
                        "sender": metadata.get("absender"),
                        "importance": metadata.get("wichtigkeit", "mittel"),
                        "deadlines": metadata.get("fristen", []),
                        "metadata": metadata,
                        "updated_at": now
                    }
                    if metadata.get("datum") and metadata["datum"] != "null":
                        update_data["document_date"] = metadata["datum"]
                    await db.documents.update_one({"id": doc["id"]}, {"$set": update_data})
                else:
                    await db.documents.update_one(
                        {"id": doc["id"]},
                        {"$set": {"ocr_text": ocr_text, "ocr_processed": True, "updated_at": now}}
                    )
                processed += 1
            else:
                errors += 1
        except Exception as e:
            logger.error(f"Batch reprocess error for {doc['id']}: {e}")
            errors += 1

    return {"success": True, "processed": processed, "errors": errors, "total": len(unprocessed)}



@router.post("/documents/assign-case")
async def assign_documents_to_case(
    document_ids: str = Form(...),
    case_id: str = Form(...),
    user: dict = Depends(require_auth)
):
    doc_ids = json.loads(document_ids)
    
    case = await db.cases.find_one({"id": case_id, "user_id": user["id"]})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    updated = 0
    for doc_id in doc_ids:
        result = await db.documents.update_one(
            {"id": doc_id, "user_id": user["id"]},
            {"$set": {"case_id": case_id, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        if result.modified_count:
            await db.cases.update_one({"id": case_id}, {"$addToSet": {"document_ids": doc_id}})
            updated += 1
    
    await log_action(user["id"], "assign_documents", "case", case_id, {"count": updated})
    return {"success": True, "updated": updated}


@router.get("/documents/suggest-for-case/{case_id}")
async def suggest_documents_for_case(
    case_id: str,
    user: dict = Depends(require_auth)
):
    """KI-basierte Suche nach passenden Dokumenten für einen Fall"""
    case = await db.cases.find_one({"id": case_id, "user_id": user["id"]})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get all user's documents that are NOT in this case
    all_docs = await db.documents.find({
        "user_id": user["id"],
        "$or": [{"case_id": {"$exists": False}}, {"case_id": ""}, {"case_id": None}]
    }).to_list(500)
    
    if not all_docs:
        # Get AI settings for consistent response structure
        settings = await db.system_settings.find_one({}, {"_id": 0})
        ai_provider = settings.get("ai_provider", "ollama") if settings else "ollama"
        openai_key = settings.get("openai_api_key") if settings else None
        
        return {
            "suggestions": [], 
            "total_available": 0,
            "ai_powered": ai_provider == "openai" and openai_key is not None,
            "message": "Keine weiteren Dokumente verfügbar"
        }
    
    # Get case documents for context
    case_docs = await db.documents.find({"case_id": case_id}).to_list(100)
    
    # Build case context
    case_context = f"Fall: {case.get('title', 'Unbekannt')}\n"
    case_context += f"Beschreibung: {case.get('description', '')}\n"
    case_context += f"Typ: {case.get('type', '')}\n"
    
    if case_docs:
        case_context += "\nBereits verknüpfte Dokumente:\n"
        for doc in case_docs[:5]:
            case_context += f"- {doc.get('display_name', doc.get('original_filename', ''))} "
            if doc.get('sender'):
                case_context += f"(Von: {doc['sender']}) "
            if doc.get('ai_summary'):
                case_context += f"- {doc['ai_summary'][:100]}"
            case_context += "\n"
    
    # Try to use AI for smart matching
    settings = await db.system_settings.find_one({}, {"_id": 0})
    ai_provider = settings.get("ai_provider", "ollama") if settings else "ollama"
    openai_key = settings.get("openai_api_key") if settings else None
    
    suggestions = []
    
    if ai_provider == "openai" and openai_key:
        try:
            import openai
            client = openai.OpenAI(api_key=openai_key)
            
            # Build document list for AI
            doc_list = []
            for i, doc in enumerate(all_docs[:50]):  # Limit to 50 docs
                doc_info = {
                    "index": i,
                    "id": doc["id"],
                    "name": doc.get("display_name", doc.get("original_filename", "")),
                    "sender": doc.get("sender", ""),
                    "date": doc.get("document_date", ""),
                    "summary": doc.get("ai_summary", "")[:200] if doc.get("ai_summary") else "",
                    "keywords": doc.get("keywords", [])[:5]
                }
                doc_list.append(doc_info)
            
            prompt = f"""Analysiere den folgenden Fall und die verfügbaren Dokumente. 
Identifiziere welche Dokumente zu diesem Fall passen könnten.

FALL-KONTEXT:
{case_context}

VERFÜGBARE DOKUMENTE:
{json.dumps(doc_list, ensure_ascii=False, indent=2)}

Antworte NUR mit einem JSON-Array der passenden Dokument-IDs und einer kurzen Begründung.
Format: [{{"id": "...", "reason": "..."}}]
Wähle maximal 10 Dokumente aus, die am relevantesten sind.
Wenn keine passenden Dokumente gefunden werden, antworte mit: []"""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Du bist ein Dokumenten-Analyst für Rechtsfälle. Antworte nur mit validem JSON."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1500,
                temperature=0.3
            )
            
            ai_response = response.choices[0].message.content.strip()
            # Clean up response
            if ai_response.startswith("```"):
                ai_response = ai_response.split("\n", 1)[1].rsplit("```", 1)[0]
            
            suggested_ids = json.loads(ai_response)
            
            for suggestion in suggested_ids:
                doc = next((d for d in all_docs if d["id"] == suggestion.get("id")), None)
                if doc:
                    suggestions.append({
                        "id": doc["id"],
                        "display_name": doc.get("display_name", doc.get("original_filename", "")),
                        "sender": doc.get("sender"),
                        "document_date": doc.get("document_date"),
                        "ai_summary": doc.get("ai_summary"),
                        "reason": suggestion.get("reason", "Möglicherweise relevant")
                    })
                    
        except Exception as e:
            logger.error(f"AI document suggestion error: {e}")
            # Fallback to keyword matching
    
    # Fallback: Simple keyword matching if AI fails or not available
    if not suggestions:
        case_keywords = case.get("title", "").lower().split() + case.get("description", "").lower().split()
        case_keywords = [k for k in case_keywords if len(k) > 3]
        
        for doc in all_docs:
            score = 0
            doc_text = (
                (doc.get("display_name", "") or "") + " " +
                (doc.get("ai_summary", "") or "") + " " +
                (doc.get("sender", "") or "") + " " +
                " ".join(doc.get("keywords", []) or [])
            ).lower()
            
            for keyword in case_keywords:
                if keyword in doc_text:
                    score += 1
            
            if score > 0:
                suggestions.append({
                    "id": doc["id"],
                    "display_name": doc.get("display_name", doc.get("original_filename", "")),
                    "sender": doc.get("sender"),
                    "document_date": doc.get("document_date"),
                    "ai_summary": doc.get("ai_summary"),
                    "reason": f"Übereinstimmende Begriffe: {score}",
                    "score": score
                })
        
        # Sort by score
        suggestions.sort(key=lambda x: x.get("score", 0), reverse=True)
        suggestions = suggestions[:10]
    
    return {
        "suggestions": suggestions,
        "total_available": len(all_docs),
        "ai_powered": ai_provider == "openai" and openai_key is not None
    }


@router.post("/documents/generate-word")
async def generate_word_document(
    title: str = Form(...),
    content: str = Form(...),
    template: str = Form("letter"),  # letter, report, contract
    recipient_name: str = Form(None),
    recipient_address: str = Form(None),
    sender_name: str = Form(None),
    sender_address: str = Form(None),
    date_str: str = Form(None),
    subject: str = Form(None),
    case_id: str = Form(None),
    user: dict = Depends(require_auth)
):
    """Generate a Word document from content
    
    Templates:
    - letter: Formal letter with sender, recipient, date, subject
    - report: Simple document with title and content
    - contract: Document with parties and terms
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    if template == "letter":
        # === LETTER TEMPLATE ===
        
        # Sender (top right)
        if sender_name or sender_address:
            sender_para = doc.add_paragraph()
            sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if sender_name:
                sender_para.add_run(sender_name + "\n")
            if sender_address:
                for line in sender_address.split("\n"):
                    sender_para.add_run(line + "\n")
            doc.add_paragraph()
        
        # Recipient (left)
        if recipient_name or recipient_address:
            recipient_para = doc.add_paragraph()
            if recipient_name:
                recipient_para.add_run(recipient_name + "\n")
            if recipient_address:
                for line in recipient_address.split("\n"):
                    recipient_para.add_run(line + "\n")
            doc.add_paragraph()
        
        # Date (right)
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(date_str or datetime.now().strftime("%d.%m.%Y"))
        doc.add_paragraph()
        
        # Subject (bold)
        if subject:
            subject_para = doc.add_paragraph()
            subject_run = subject_para.add_run(subject)
            subject_run.bold = True
            doc.add_paragraph()
        
        # Content
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Signature area
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("Mit freundlichen Grüßen")
        doc.add_paragraph()
        doc.add_paragraph()
        if sender_name:
            doc.add_paragraph(sender_name)
    
    elif template == "report":
        # === REPORT TEMPLATE ===
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(date_str or datetime.now().strftime("%d.%m.%Y"))
        doc.add_paragraph()
        
        # Content
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    
    else:  # contract or default
        # === CONTRACT/DEFAULT TEMPLATE ===
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Parties
        if sender_name and recipient_name:
            parties_para = doc.add_paragraph()
            parties_para.add_run("Zwischen:\n").bold = True
            parties_para.add_run(f"{sender_name}\n")
            if sender_address:
                parties_para.add_run(f"{sender_address}\n")
            parties_para.add_run("\nund\n\n").bold = True
            parties_para.add_run(f"{recipient_name}\n")
            if recipient_address:
                parties_para.add_run(f"{recipient_address}\n")
            doc.add_paragraph()
        
        # Content
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Signature area
        doc.add_paragraph()
        doc.add_paragraph()
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.cell(0, 0).text = "Ort, Datum"
        sig_table.cell(0, 1).text = "Ort, Datum"
        sig_table.cell(2, 0).text = sender_name or "_______________"
        sig_table.cell(2, 1).text = recipient_name or "_______________"
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Save to file system
    doc_id = str(uuid.uuid4())
    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:50]
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"
    storage_filename = f"{doc_id}.docx"
    storage_path = os.path.join(UPLOAD_DIR, storage_filename)
    
    async with aiofiles.open(storage_path, 'wb') as f:
        await f.write(doc_buffer.getvalue())
    
    # Create document record
    document = {
        "id": doc_id,
        "user_id": user["id"],
        "original_filename": filename,
        "display_name": filename,
        "storage_path": storage_path,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file_size": len(doc_buffer.getvalue()),
        "document_type": template,
        "ocr_text": content,  # Store the content as OCR text for searchability
        "ocr_processed": True,
        "ai_analyzed": False,
        "tags": ["generiert", template],
        "case_id": case_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.documents.insert_one(document)
    await log_action(user["id"], "generate_document", "document", doc_id, {"filename": filename, "template": template})
    
    # If case_id, link to case
    if case_id:
        await db.cases.update_one(
            {"id": case_id},
            {"$addToSet": {"document_ids": doc_id}}
        )
    
    return {
        "success": True,
        "document_id": doc_id,
        "filename": filename,
        "message": f"Dokument '{filename}' wurde erstellt"
    }
